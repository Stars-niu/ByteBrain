# 导入所需的库
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch
import streamlit as st
from modelscope import snapshot_download
from transformers import DPRQuestionEncoder, DPRQuestionEncoderTokenizer, DPRContextEncoder, DPRContextEncoderTokenizer
import fitz  # PyMuPDF
import docx
import faiss
import numpy as np

# 设置页面配置
st.set_page_config(
    page_title="ByteBrain",
    page_icon="🪐",
    layout="wide",
    initial_sidebar_state="auto",
    menu_items={
        "Get Help": "https://bytebrain.com/help",
        "Report a bug": "https://bytebrain.com/bug-report",
        "About": "https://bytebrain.com/about"
    }
)

# 添加自定义CSS样式
st.markdown(
    """
    <style>
    .main {
        background-image: url('Background.png');
        background-size: cover;
        background-position: center;
        padding: 10px;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: blue;
        border: none;
        padding: 10px 24px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        margin: 4px 2px;
        cursor: pointer;
        border-radius: 12px;
    }
    .stTextInput>div>div>input {
        border: 2px solid #4CAF50;
        border-radius: 12px;
        padding: 10px;
    }
    .fixed-right {
        position: fixed;
        top: 20px;
        right: 20px;
        width: 30%;
    }
    .chat-container {
        width: 100%;
        max-height: 70vh;
        overflow-y: auto;
        padding-right: 20px;
    }
    .stChatMessage {
        width: 100%;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    '<span style="font-size: 60px">✨ ByteBrain</span>&nbsp;&nbsp;<span style="font-size: 24px">——计算机科学智能知识助手</span>',
    unsafe_allow_html=True
)

# 源大模型下载
model_dir = snapshot_download('IEITYuan/Yuan2-2B-July-hf', cache_dir='./')

# 定义模型路径
path = './IEITYuan/Yuan2-2B-July-hf'

# 定义模型数据类型
torch_dtype = torch.bfloat16 # A10

# 定义一个函数，用于获取模型和tokenizer
@st.cache_resource
def get_model():
    print("创建tokenizer...")
    tokenizer = AutoTokenizer.from_pretrained(path, add_eos_token=False, add_bos_token=False, eos_token='<eod>')
    tokenizer.add_tokens(['<sep>', '<pad>', '<mask>', '<predict>', '<FIM_SUFFIX>', '<FIM_PREFIX>', '<FIM_MIDDLE>','<commit_before>','<commit_msg>','<commit_after>','<jupyter_start>','<jupyter_text>','<jupyter_code>','<jupyter_output>','<empty_output>'], special_tokens=True)

    print("创建模型...")
    model = AutoModelForCausalLM.from_pretrained(path, torch_dtype=torch_dtype, trust_remote_code=True).cuda()

    print("完成.")
    return tokenizer, model

# 加载model和tokenizer
tokenizer, model = get_model()

# 加载DPR模型和tokenizer
@st.cache_resource
def get_dpr_models():
    question_encoder = DPRQuestionEncoder.from_pretrained("facebook/dpr-question_encoder-single-nq").cuda()
    question_tokenizer = DPRQuestionEncoderTokenizer.from_pretrained("facebook/dpr-question_encoder-single-nq")
    context_encoder = DPRContextEncoder.from_pretrained("facebook/dpr-ctx_encoder-single-nq").cuda()
    context_tokenizer = DPRContextEncoderTokenizer.from_pretrained("facebook/dpr-ctx_encoder-single-nq")
    return question_encoder, question_tokenizer, context_encoder, context_tokenizer

question_encoder, question_tokenizer, context_encoder, context_tokenizer = get_dpr_models()

# 读取PDF文件
def read_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# 读取Word文件
def read_word(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# 加载知识库
@st.cache_resource
def load_knowledge_base():
    return []

knowledge_base = load_knowledge_base()

# 构建向量数据库
@st.cache_resource
def build_faiss_index(knowledge_base):
    context_embeddings = []
    for doc in knowledge_base:
        inputs = context_tokenizer(doc, return_tensors="pt", truncation=True, padding=True).to("cuda")
        embedding = context_encoder(**inputs).pooler_output.cpu().detach().numpy()
        context_embeddings.append(embedding)
    context_embeddings = np.vstack(context_embeddings)
    index = faiss.IndexFlatL2(context_embeddings.shape[1])
    index.add(context_embeddings)
    return index

index = build_faiss_index(knowledge_base)

# 初次运行时，session_state中没有"messages"，需要创建一个空列表
if "messages" not in st.session_state:
    st.session_state["messages"] = []

# 使用分栏布局
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("<div class='chat-container'>", unsafe_allow_html=True)
    # 每次对话时，都需要遍历session_state中的所有消息，并显示在聊天界面上
    for msg in st.session_state.messages:
        st.chat_message(msg["role"]).write(msg["content"])
    st.markdown("</div>", unsafe_allow_html=True)

with col2:
    st.markdown("<div class='fixed-right'>", unsafe_allow_html=True)
    st.image("logo.png", caption="ByteBrain Logo", width=150)
    st.markdown("ByteBrain——一个智能知识助手，旨在帮助用户快速获取信息和解决问题。")
    st.markdown("### 联系我们")
    st.markdown("如果您有任何问题或建议，请通过以下方式联系我们：")
    st.markdown("- 邮箱: support@bytebrain.com")
    st.markdown("- 电话: 520-1314")
    st.markdown("</div>", unsafe_allow_html=True)

# 文件上传
uploaded_file = st.file_uploader("上传PDF或Word文件", type=["pdf", "docx"])

if uploaded_file:
    if uploaded_file.type == "application/pdf":
        file_text = read_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        file_text = read_word(uploaded_file)
    
    # 将文件内容添加到知识库中
    knowledge_base.append(file_text)
    
    # 重新构建向量数据库
    index = build_faiss_index(knowledge_base)

# 聊天输入框放在col1之外
prompt = st.chat_input("请输入您的问题:")

if prompt:
    # 将用户的输入添加到session_state中的messages列表中
    st.session_state.messages.append({"role": "user", "content": prompt})

    # 在聊天界面上显示用户的输入
    st.chat_message("user").write(prompt)

    # 使用DPR模型进行检索
    inputs = question_tokenizer(prompt, return_tensors="pt", truncation=True, padding=True).to("cuda")
    question_embedding = question_encoder(**inputs).pooler_output.cpu().detach().numpy()
    D, I = index.search(question_embedding, k=5)  # 检索前5个相关文档
    retrieved_docs = [knowledge_base[i] for i in I[0]]

    # 将检索到的文档拼接到prompt中
    prompt_with_docs = prompt + "\n\n" + "\n\n".join(retrieved_docs) + "<sep>"

    # 调用生成模型
    inputs = tokenizer(prompt_with_docs, return_tensors="pt")["input_ids"].cuda()
    outputs = model.generate(inputs, do_sample=False, max_length=1024)  # 设置解码方式和最大生成长度
    output = tokenizer.decode(outputs[0])
    response = output.split("<sep>")[-1].replace("<eod>", '')

    # 将模型的输出添加到session_state中的messages列表中
    st.session_state.messages.append({"role": "assistant", "content": response})

    # 在聊天界面上显示模型的输出
    st.chat_message("assistant").write(response)
# 代码说明
# 读取PDF和Word文件：使用fitz库读取PDF文件，使用docx库读取Word文件。
# 文件上传功能：使用Streamlit的file_uploader组件允许用户上传PDF或Word文件。
# 将文件内容添加到知识库：读取文件内容并添加到知识库中。
# 重新构建向量数据库：每次添加新文件后，重新构建向量数据库。

# 可以使用一些Python库来处理PDF和Word文档，例如PyMuPDF（又名fitz）用于PDF，python-docx用于Word文档。

# 下面是一个示例，展示了如何读取PDF和Word文档并将其内容添加到知识库中：

# 安装所需的库
# bash
# pip install PyMuPDF python-docx
# 读取PDF和Word文档的示例代码
# python
# import fitz  # PyMuPDF
# import docx

# def read_pdf(file_path):
#     """读取PDF文件并返回文本内容"""
#     doc = fitz.open(file_path)
#     text = ""
#     for page in doc:
#         text += page.get_text()
#     return text

# def read_word(file_path):
#     """读取Word文件并返回文本内容"""
#     doc = docx.Document(file_path)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

# # 示例：读取PDF和Word文件
# pdf_text = read_pdf("example.pdf")
# word_text = read_word("example.docx")

# # 将读取的文本添加到知识库中
# knowledge_base = [pdf_text, word_text]
# 将读取PDF和Word文档的功能集成到现有的Streamlit应用中
# 你可以在Streamlit应用中添加文件上传功能，并在用户上传文件后读取其内容并添加到知识库中。
